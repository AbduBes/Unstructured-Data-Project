from docx import Document
import chardet
import os
from datetime import datetime

import sys
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', 'src')))

from storage.mongo import save_to_mongo
from utils.logger import logging


def detect_encoding(file_path):
    """Detect file encoding using chardet."""
    with open(file_path, "rb") as f:
        raw = f.read()
    result = chardet.detect(raw)
    encoding = result.get("encoding", "utf-8") or "utf-8"
    logging.info(f"Detected encoding for {file_path}: {encoding}")
    return encoding


def extract_text_from_word(file_path):
    """Extract paragraphs, runs, tables, and styles from a Word document."""
    extracted = {
        "paragraphs": [],
        "tables": []
    }

    try:
        doc = Document(file_path)

        # Extract paragraphs with style info
        for para in doc.paragraphs:
            if para.text.strip():
                runs_text = [run.text for run in para.runs if run.text.strip()]
                extracted["paragraphs"].append({
                    "text": para.text.strip(),
                    "style": para.style.name,
                    "runs": runs_text
                })

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            extracted["tables"].append({
                "table_index": table_idx,
                "rows": table_data
            })

        logging.info(f"Extracted {len(extracted['paragraphs'])} paragraphs and "
                     f"{len(extracted['tables'])} tables from: {file_path}")

    except Exception as e:
        logging.error(f"Failed to extract Word document {file_path}: {e}")

    return extracted


def process_word(file_path):
    """Extract, structure, and store Word document content in MongoDB."""
    file_name = os.path.basename(file_path)
    logging.info(f"Processing Word document: {file_name}")

    content = extract_text_from_word(file_path)

    document = {
        "content": content,
        "metadata": {
            "file_name": file_name,
            "document_type": "word",
            "extraction_timestamp": datetime.utcnow(),
            "source": file_path,
            "extraction_library": "python-docx"
        }
    }

    save_to_mongo(document, source_url=file_path)
    logging.info(f"Stored Word document {file_name} in MongoDB")


def process_all_word_docs(folder_path="data/raw/word"):
    """Process all .docx files in a folder."""
    if not os.path.exists(folder_path):
        logging.warning(f"Word folder not found: {folder_path}")
        return

    word_files = [f for f in os.listdir(folder_path) if f.endswith(".docx")]
    if not word_files:
        logging.warning(f"No Word files found in {folder_path}")
        return

    for filename in word_files:
        file_path = os.path.join(folder_path, filename)
        process_word(file_path)


if __name__ == "__main__":
    process_all_word_docs()